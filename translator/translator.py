import subprocess
import re
import os
import sys
from pathlib import Path
import google.generativeai as genai
from collections import Counter

# 尝试导入docx库，如果失败则给出安装提示
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("\n[错误] 未找到 'python-docx' 库。")
    print("请先运行 'pip install python-docx' 来安装它。")
    sys.exit(1)

# =================== 路径配置 ===================
BASE_DIR = Path('/workspaces/xiaoqizhangxz-arch.github.io/translator')
INPUT_PDF_FOLDER = BASE_DIR / 'source_pdfs'
OUTPUT_TXT_FOLDER = BASE_DIR / 'cleaned_txts'
OUTPUT_TRANSLATION_FOLDER = BASE_DIR / 'translated_zh_docx' # 输出文件夹改为docx

# =================== Gemini 配置 ===================
MODEL_NAME = "models/gemini-2.5-pro"

TRANSLATION_PROMPT_TEMPLATE = """
你是一位顶级的学术翻译专家，专攻将英文的人文社科类学术著作翻译成流畅、精准、且符合学术规范的简体中文。

你的任务是翻译以下完整的英文书籍文本。

请严格遵循以下准则：
1.  **忠实原文**: 精确传达原文的复杂概念、论点和细微语气。不得添加任何原文没有的解释或评论。
2.  **术语一致**: 确保关键的理论术语和专有名词在全文中保持翻译的一致性。
3.  **行文流畅**: 译文必须通顺、自然，符合中文学术写作的语言习惯。
4.  **格式保留**: 保持原文的段落结构。原文中的换行和分段应在译文中得到保留。
5.  **纯净输出**: 你的回答必须是且仅是翻译后的简体中文文本。不要包含任何“好的，这是翻译：”或任何其他前言或结语。

--- English Text to Translate ---
{full_text}
---
"""

# =================== 功能函数 ===================

def extract_text_with_pdftotext(pdf_path):
    """使用pdftotext从PDF中提取高质量的、保持布局的文本。"""
    print(f"   -> 正在使用 pdftotext 提取文本...")
    try:
        result = subprocess.run(
            ['pdftotext', '-layout', '-enc', 'UTF-8', str(pdf_path), '-'],
            capture_output=True, text=True, check=True, encoding='utf-8'
        )
        return result.stdout
    except FileNotFoundError:
        print("\n[错误] 未找到 'pdftotext' 命令。请运行 'sudo apt-get update && sudo apt-get install -y poppler-utils'。")
        sys.exit(1)
    except subprocess.CalledProcessError as e:
        print(f"\n[警告] 使用pdftotext处理文件 '{pdf_path.name}' 时出错: {e.stderr}")
        return None

def clean_book_text(raw_text):
    """增强版清理函数，能更有效地移除页眉、页脚和页码。"""
    print("   -> 正在智能清理文本...")
    if not raw_text: return ""
    
    # 初步清理：修复断词、移除分页符
    text = re.sub(r'-\s*\n\s*', '', raw_text)
    text = text.replace('\f', '')
    
    lines = text.split('\n')
    
    # 统计所有行的出现频率，以识别重复的页眉/页脚
    line_counts = Counter(line.strip() for line in lines if len(line.strip()) > 5)
    # 找出出现超过5次且长度小于100的行，作为页眉/页脚嫌疑对象
    common_headers_footers = {line for line, count in line_counts.items() if count > 5 and len(line) < 100}

    cleaned_lines = []
    for line in lines:
        stripped_line = line.strip()
        
        # 规则1: 跳过空行
        if not stripped_line:
            cleaned_lines.append("") # 保留空行以维持段落结构
            continue
        
        # 规则2: 跳过已识别的重复页眉/页脚
        if stripped_line in common_headers_footers:
            continue
            
        # 规则3: 跳过看起来像页码的行（纯数字或 "Page X" 等）
        if stripped_line.isdigit() or re.match(r'^[Pp]age\s*\d+$', stripped_line):
            continue
            
        cleaned_lines.append(line)
        
    text = '\n'.join(cleaned_lines)
    # 压缩多个连续的空行到一个，保持段落结构
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

def save_as_docx(text, file_path):
    """将文本保存为带格式的DOCX文件。"""
    print(f"   -> 正在生成格式化的 DOCX 文件...")
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Heading 1'].font.name = '宋体'
    doc.styles['Heading 1']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 设置正文样式：小四，首行缩进2字符
    style = doc.styles['Normal']
    style.font.size = Pt(12) # 小四对应12号字
    p_format = style.paragraph_format
    p_format.first_line_indent = Inches(0.33) # 约等于2个字符
    p_format.line_spacing = 1.5

    # 设置一级标题样式：四号，不缩进
    style_h1 = doc.styles['Heading 1']
    style_h1.font.size = Pt(14) # 四号对应14号字
    style_h1.font.bold = True
    p_format_h1 = style_h1.paragraph_format
    p_format_h1.first_line_indent = Inches(0)
    p_format_h1.space_before = Pt(12)
    p_format_h1.space_after = Pt(12)
    
    paragraphs = text.split('\n')
    for para_text in paragraphs:
        stripped_para = para_text.strip()
        if not stripped_para:
            continue
        
        # 启发式规则判断章节标题：行长较短（小于50字符），且不以句号、问号等结尾
        is_heading = len(stripped_para) < 50 and not stripped_para.endswith(('.', '?', '!', '”'))
        
        if is_heading:
            doc.add_heading(stripped_para, level=1)
        else:
            doc.add_paragraph(para_text)
            
    doc.save(file_path)

def translate_text_with_gemini(cleaned_text):
    """使用Gemini API翻译完整的文本。"""
    api_key = os.environ.get('GEMINI_API_KEY')
    if not api_key:
        print("\n[错误] Gemini API 密钥未在环境变量 GEMINI_API_KEY 中设置。")
        return None
    
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(MODEL_NAME)
    
    full_prompt = TRANSLATION_PROMPT_TEMPLATE.format(full_text=cleaned_text)
    
    try:
        print("   -> 已发送完整文本至 Gemini API，等待翻译... (这对于一本书来说可能需要很长时间)")
        response = model.generate_content(full_prompt, request_options={'timeout': 3600}) # 1小时超时
        
        if hasattr(response, 'text'):
            return response.text
        if hasattr(response, 'candidates') and response.candidates:
            if hasattr(response.candidates[0], 'content') and response.candidates[0].content.parts:
                return response.candidates[0].content.parts[0].text
        return str(response)

    except Exception as e:
        print(f"\n[错误] Gemini API 请求阶段出错: {e}")
        return None

# =================== 主流程 ===================
def main():
    INPUT_PDF_FOLDER.mkdir(exist_ok=True)
    OUTPUT_TXT_FOLDER.mkdir(exist_ok=True)
    OUTPUT_TRANSLATION_FOLDER.mkdir(exist_ok=True)

    pdf_files = [f for f in INPUT_PDF_FOLDER.iterdir() if f.suffix.lower() == '.pdf' and not f.name.startswith('._')]
    if not pdf_files:
        print(f"[信息] 在 '{INPUT_PDF_FOLDER}' 文件夹中没有找到PDF文件。")
        return
    
    total_files = len(pdf_files)
    print(f"找到 {total_files} 本书 (PDF)，开始翻译流程...")

    for i, pdf_path in enumerate(pdf_files):
        print(f"\n--- [{i+1}/{total_files}] 正在处理: {pdf_path.name} ---")
        base_name = pdf_path.stem
        
        txt_path = OUTPUT_TXT_FOLDER / f"{base_name}_cleaned.txt"
        translation_path = OUTPUT_TRANSLATION_FOLDER / f"{base_name}_translated_zh.docx" # 输出文件后缀改为.docx
        
        print("1. 清理PDF为TXT...")
        raw_text = extract_text_with_pdftotext(pdf_path)
        if not raw_text:
            print(f"   跳过，文本提取失败。")
            continue
        
        cleaned_text = clean_book_text(raw_text)
        txt_path.write_text(cleaned_text, encoding='utf-8')
        print(f"   清理后的TXT已保存至: {txt_path.name}")

        print("2. 翻译完整的TXT文件...")
        translated_text = translate_text_with_gemini(cleaned_text)
        
        if translated_text and len(translated_text) > 100:
            # 使用新函数保存为带格式的DOCX
            save_as_docx(translated_text, translation_path)
            print(f"   ✅ 翻译完成！已生成格式化DOCX文件: {translation_path.name}")
        else:
            print(f"   ❌ 未能为 {pdf_path.name} 生成有效的翻译。")

    print(f"\n--- 所有任务完成 ---")
    print(f"请在以下文件夹中查看结果:")
    print(f"- 清理后的TXT文件: {OUTPUT_TXT_FOLDER}")
    print(f"- 中文译文 (DOCX): {OUTPUT_TRANSLATION_FOLDER}")

if __name__ == '__main__':
    main()